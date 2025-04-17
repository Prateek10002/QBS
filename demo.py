import streamlit as st
import openai
import base64
import os
from dotenv import load_dotenv
from textstat import flesch_reading_ease
from PIL import Image, UnidentifiedImageError
import io
import pypdf
from transformers import pipeline, AutoModelForCausalLM, AutoTokenizer
import torch
# from gpt4all import GPT4All

from dotenv import load_dotenv
import os
load_dotenv()
# Load environment variables explicitly
dotenv_path = os.path.join(os.path.dirname(__file__), ".env")
load_dotenv(dotenv_path)


def get_base64_of_bin_file(bin_file):
    if not os.path.exists(bin_file):
        return ""  
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

img_file = "C:\\Users\\prateek.kumar\\Desktop\\Langchain Project\\QBS (4).jpg"

base64_img = get_base64_of_bin_file(img_file)

if base64_img:
    page_bg_img = f'''
    <style> 
    .stApp {{
        background-image: url("data:image/png;base64,{base64_img}");
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
    }}
    label {{
        color: black;
    }}
    .response-box {{
        background-color: rgba(255, 255, 255, 0.9);
        padding: 20px;
        border-radius: 10px;
        color: black;
        margin-top: 10px;
    }} 
    h1 {{
        color: black;
        text-align: center;
    }}
    .content-box {{
        background-color: rgba(0, 0, 0, 0.7);
        padding: 15px;
        border-radius: 10px;
        color: black;
        text-align: center;
    }}
    </style>
    '''
    st.markdown(page_bg_img, unsafe_allow_html=True)
else:
    st.warning("Background image not found. Proceeding without a background image.")

st.sidebar.title("Settings")
st.markdown("""
    <style>
    button[kind="icon"] {
        display: none !important;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown(
    """
    <style>
        /* Force all sidebar labels to be black */
        section[data-testid="stSidebar"] label {
            color: black !important;
        }

        /* Force all sidebar text elements to be black */
        section[data-testid="stSidebar"] div, 
        section[data-testid="stSidebar"] h1, 
        section[data-testid="stSidebar"] h2, 
        section[data-testid="stSidebar"] h3, 
        section[data-testid="stSidebar"] span {
            color: black !important;
        }

        /* Ensure dropdowns, sliders, and inputs have black text */
        section[data-testid="stSidebar"] select,
        section[data-testid="stSidebar"] input,
        section[data-testid="stSidebar"] div[role="slider"] {
            color: black !important;
        }
    </style>
    """,
    unsafe_allow_html=True
)

user_role = st.sidebar.selectbox("Are you an:", ["Instructor", "Learner"])
grade_level = st.sidebar.selectbox("Select Grade Level:", ["K-2", "3-5", "6-8", "9-12"])


model_mapping = {
    "Model-1": "gpt-4",
    "Model-2": "gpt-4-turbo",
    "Model-3": "gpt-4o",
   # "OpenSource": "llama3.2-vision:latest",
    "Image": "dall-e-3"
}

# Display only the mapped names in the dropdown
llm_model_display = st.sidebar.selectbox(
    "Select an AI Model",
    list(model_mapping.keys())  # Show mapped names only
)

# Get the actual model identifier for processing
llm_model = model_mapping[llm_model_display]


temperature = st.sidebar.slider("Temperature", 0.0, 1.0, 0.7)
max_tokens = st.sidebar.slider("Max Tokens", 10, 2000, 500)

def generate_openai_response(question, model, api_key, temperature, max_tokens):
    if not api_key:
        return "Error"
    try:
        openai.api_key = api_key
        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": question}],
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return response["choices"][0]["message"]["content"]
    except Exception as e:
        return f"Error: {str(e)}"

def generate_open_source_response(question, model_name):
    try:
        if model_name == "ollama":
            model_name = "llama3.2-vision:latest"
        # elif model_name == "mistral":
        #     #model_name = "mistralai/Mistral-7B"
        # elif model_name == "falcon":
        #     #model_name = "tiiuae/falcon-7b-instruct"
        else:
            return "Error: Unsupported model."

        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForCausalLM.from_pretrained(model_name)
        device = 0 if torch.cuda.is_available() else -1
        generator = pipeline("text-generation", model=model, tokenizer=tokenizer, device=device)
        response = generator(question, max_length=max_tokens, num_return_sequences=1, temperature=temperature)
        return response[0]["generated_text"]
    except Exception as e:
        return f"Error: {str(e)}"


def generate_response(question):
    if llm_model in ["gpt-4", "gpt-3.5-turbo", "gpt-4-turbo", "gpt-4o", "gpt-4o-mini"]:
        return generate_openai_response(question, llm_model, openai.api_key, temperature, max_tokens)
    # elif llm_model == "gpt4all":
    #     return generate_gpt4all_response(question)
    else:
        return generate_open_source_response(question, llm_model)

# def grammar_check(text):
#     prompt = (
#         "Fix the grammar and improve the clarity of the following text:\n\n"
#         f"{text}\n\nReturn only the corrected version."
#     )
#     return generate_response(prompt)

def solve_math_problem(problem):
    prompt = f"Solve the following math problem step-by-step: {problem}"
    return generate_response(prompt)

# def generate_image(prompt):
#     try:
#         if llm_model == "dall-e-3":
#             response = openai.Image.create(
#                 model="dall-e-3",
#                 prompt=prompt,
#                 n=1,
#                 size="1024x1024"
#             )
#             return response['data'][0]['url']
#         else:
#             return "Error: Image generation is only supported with DALL-E 3."
#     except Exception as e:
#         return f"Error: {str(e)}"

def generate_image(prompt):
    # Initialize image cache in session state if it doesn't exist
    if "image_cache" not in st.session_state:
        st.session_state["image_cache"] = {}

    # Return cached image if prompt already used
    if prompt in st.session_state["image_cache"]:
        return st.session_state["image_cache"][prompt]

    # Generate a new image only if prompt is new
    if llm_model == "dall-e-3":
        try:
            response = openai.Image.create(
                model="dall-e-3",
                prompt=prompt,
                n=1,
                size="1024x1024"
            )
            image_url = response['data'][0]['url']

            # Save it in cache
            st.session_state["image_cache"][prompt] = image_url
            return image_url
        except Exception as e:
            return f"Error: {str(e)}"
    else:
        return "Error:"



def generate_content(prompt):
    return generate_response(prompt)

def adjust_text_for_grade(text, grade_level):
    if grade_level == "K-2":
        target_score = 90  # Simple readability
    elif grade_level == "3-5":
        target_score = 80
    elif grade_level == "6-8":
        target_score = 70
    elif grade_level == "9-12":
        target_score = 60
    else:
        target_score = 50

    while flesch_reading_ease(text) < target_score:
        text = simplify_text(text)  # Simplify text function to adjust readability
    return text

def simplify_text(text):
    return text  # Implement text simplification logic here

import io
import docx2txt
from pdf2image import convert_from_bytes
from PIL import Image
import streamlit as st
import ollama

with st.sidebar:
    st.header("Upload")
    uploaded_file = st.file_uploader("Upload a file", type=['png', 'jpg', 'jpeg', 'pdf', 'docx', 'txt'])

    if uploaded_file is not None:
        all_text = ""

        try:
            # 🖼️ IMAGE
            if uploaded_file.type.startswith("image/"):
                image = Image.open(uploaded_file)
                st.image(image, caption="Uploaded Image")

                if st.button("Extract Text 🔍", type="primary"):
                    with st.spinner("Processing image..."):
                        response = ollama.chat(
                        model='llama3.2-vision:latest',
                        messages=[{
                                    'role': 'user',
                                    'content': 'Extract all readable content and present it in structured Markdown format.',
                                    'images': [uploaded_file.getvalue()]
                                }]
                            )

                        st.session_state['ocr_result'] = response.message.content

            # 📄 PDF
            elif uploaded_file.type == "application/pdf":
                pdf_images = convert_from_bytes(uploaded_file.getvalue())
                for i, img in enumerate(pdf_images):
                    st.image(img, caption=f"PDF Page {i+1}")

                if st.button("Extract Text 🔍", type="primary"):
                    with st.spinner("Processing PDF..."):
                        all_text = ""
                        for idx, img in enumerate(pdf_images):
                            buf = io.BytesIO()
                            img.save(buf, format="PNG")
                            buf.seek(0)

                            response = ollama.chat(
                                model='llama3.2-vision:latest',
                                messages=[{
                                    'role': 'user',
                                    'content': f"""Extract text from page {idx+1} in structured Markdown format.""",
                                    'images': [buf.getvalue()]
                                }]
                            )
                            all_text += f"## Page {idx+1}\n" + response.message.content + "\n\n---\n\n"

                        st.session_state['ocr_result'] = all_text

            # 📃 DOCX
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                if st.button("Extract Text 🔍", type="primary"):
                    with st.spinner("Reading DOCX file..."):
                        raw_text = docx2txt.process(uploaded_file)
                        response = ollama.chat(
                            model='llama3.2-vision:latest',
                            messages=[{
                                'role': 'user',
                                'content': f"""Format this DOCX content clearly in Markdown:\n\n{raw_text}"""
                            }]
                        )
                        st.session_state['ocr_result'] = response.message.content

            #
            elif uploaded_file.type == "text/plain":
                if st.button("Extract Text 🔍", type="primary"):
                    with st.spinner("Reading text file..."):
                        raw_text = uploaded_file.read().decode("utf-8")
                        response = ollama.chat(
                            model='llama3.2-vision:latest',
                            messages=[{
                                'role': 'user',
                                'content': f"""Format this plain text content in Markdown:\n\n{raw_text}"""
                            }]
                        )
                        st.session_state['ocr_result'] = response.message.content

        except Exception as e:
            st.error(f"Error: {str(e)}")

# # Display extracted content, if available
# if 'ocr_result' in st.session_state and st.session_state['ocr_result']:
#     st.markdown("## 📄 Extracted Content")
#     st.markdown(st.session_state['ocr_result'])

if 'ocr_result' in st.session_state and st.session_state['ocr_result']:
    st.markdown('<h2 style="color: white;">📄 Extracted Content</h2>', unsafe_allow_html=True)
    styled_content = f"""
    <div style="
        background-color: white;
        color: black;
        padding: 20px;
        border-radius: 10px;
        overflow-x: auto;
        white-space: pre-wrap;
        font-size: 16px;
        line-height: 1.6;
        box-shadow: 0 0 10px rgba(0,0,0,0.15);
    ">
    {st.session_state['ocr_result']}
 
    """
    st.markdown(styled_content, unsafe_allow_html=True)

    st.markdown(
        '<h3 style="color: white;">Ask a question about this content:</h3>',
        unsafe_allow_html=True
    )

    user_query = st.text_input("Enter your question", key="doc_qa")

    if st.button("Ask", key="ask_doc_question"):
        if user_query.strip():
            with st.spinner("Thinking..."):
                context = st.session_state['ocr_result']
                qa_prompt = f"""Based on the document content below, Answer the user's question.

---DOCUMENT---
{context}
---QUESTION---
{user_query}

Give a helpful, clear answer. If the document doesn't contain enough info, say so."""

                response = ollama.chat(
                    model='llama3.2-vision:latest',  # or any QA-capable model
                    messages=[{"role": "user", "content": qa_prompt}]
                )

                st.markdown(
                    '<h3 style="color: white;"> Response:</h3>',
                    unsafe_allow_html=True
                )

                st.markdown(
                    f"<div style='background-color:#f0f0f0; padding:15px; border-radius:10px; color:black;'>{response['message']['content']}</div>",
                    unsafe_allow_html=True
                )
        else:
            st.warning("Please type a question to ask.")


st.markdown(
    """
    <style>
        /* Move the welcome text down */
        .welcome-text {
            color: white;
            text-align: center;
            font-size: 36px;
            font-weight: bold;
            margin-top: 80px; /* Adjust this value to move it further down */
        }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown(f'<h1 class="welcome-text">Welcome {user_role}!</h1>', unsafe_allow_html=True)

#-----------------------------Working-------------------
st.markdown(
    """
    <style>
    label {
        color: black !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

user_prompt = st.text_area("Enter your input:")

st.markdown(
    """
    <style>
    label {
        color: white !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)
#------------------Working----------------------------
task_type = st.selectbox("Select Task:", ["Generate Content", "Solve Math Problem", "Generate Image"])
#--------------------------------------------------------

if st.button("Process Task"):
    st.session_state.pop("ocr_result", None)
    st.session_state.pop("doc_qa", None)
    if user_prompt:
        # if task_type == "Grammar Check":
        #     result = grammar_check(user_prompt)
        #     st.markdown("### Grammar Check Result:")
        #     st.markdown(f"<div class='response-box'>{result}</div>", unsafe_allow_html=True)
        if task_type == "Solve Math Problem":
            result = solve_math_problem(user_prompt)
            st.markdown("### Math Solution:")
            st.markdown(f"<div class='response-box'>{result}</div>", unsafe_allow_html=True)
        elif task_type == "Generate Image":
            result = generate_image(user_prompt)
            if "Error" not in result:
                st.image(result, caption="Generated Image")
            else:
                st.error(result)
        elif task_type == "Generate Content":
            result = generate_content(user_prompt)
            st.markdown('<h3 style="color: white;">Generated Content:</h3>', unsafe_allow_html=True)
            st.markdown(f"<div class='response-box'>{result}</div>", unsafe_allow_html=True)
    else:
        st.warning("Please provide input for processing.")

from PyPDF2 import PdfReader
import ollama

# st.sidebar.markdown("---")
# st.sidebar.subheader("📝 AI Proofreader")

# # ---- FILE UPLOAD ----
# st.sidebar.markdown("📎 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# proof_input = st.sidebar.text_area(
#     "Enter text to improve:",
#     height=150,
#     key="proof_input"
# )

# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"])
# style = st.sidebar.selectbox("Style", ["Paraphrase", "Simplify", "Make Concise", "Expand"])
# ollama_model = "llama3.2-vision:latest"  

# # ---- IMPROVE BUTTON ----
# if st.sidebar.button("Improve Text"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)
    
#     if proof_input.strip():
#         with st.spinner("Improving your text using AI Proofreader..."):
#             ...
#             system_prompt = f"""
# You are a helpful writing assistant. Improve the following text while maintaining its original meaning.
# - Tone: {tone}
# - Style: {style}

# Focus on grammar, clarity, and fluency.
# Return only the improved version of the text.
# """

#             try:
#                 response = ollama.chat(
#                     model=ollama_model,
#                     messages=[
#                         {"role": "system", "content": system_prompt},
#                         {"role": "user", "content": proof_input}
#                     ]
#                 )
#                 improved_text = response['message']['content']

#                 # st.sidebar.markdown("**Improved Text:**")
#                 # st.sidebar.markdown(
#                 #     f"<div style='background-color:#f0f0f0; padding:10px; border-radius:8px; color: black;'>{improved_text}</div>",
#                 #     unsafe_allow_html=True
#                 st.session_state["proof_result"] = improved_text

                

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text.")

# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;">📝 Improved Text:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

# # ----- Extra Writing Tools -----
# st.sidebar.markdown("### ✨ More AI Writing Tools")
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# tool_input = st.sidebar.text_area("Enter text for the selected tool:", key="tool_input")

# if st.sidebar.button("Run Tool"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     if tool_input.strip():
#         tool_prompt_map = {
#             "Paraphraser": "Paraphrase the following text while keeping its meaning.",
#             "Grammar Checker": "Fix the grammar in the following text.",
#             "AI Detector": "Analyze whether this text appears to be written by an AI or a human. Explain your reasoning.",
#             "Plagiarism Checker": "Check this text for potential plagiarism. If found, highlight the copied parts.",
#             "Summarizer": "Summarize the following text clearly and concisely."
#         }

#         prompt = tool_prompt_map.get(tool_option, "Process this text:")

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 response = ollama.chat(
#                     model=ollama_model,
#                     messages=[
#                         {"role": "system", "content": prompt},
#                         {"role": "user", "content": tool_input}
#                     ]
#                 )
#                 tool_output = response['message']['content']

#                 st.markdown(f'<h3 style="color: white;">🛠️ {tool_option} Result:</h3>', unsafe_allow_html=True)
#                 st.markdown(
#                     f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{tool_output}</div>",
#                     unsafe_allow_html=True
#                 )

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text for processing.")

# from PyPDF2 import PdfReader
# import ollama

# -------------- Sidebar: AI Proofreader with Tool Selection ---------------
# st.sidebar.markdown("---")
# st.sidebar.subheader("📝 AI Proofreader")

# # ✨ 1. Choose Tool First
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # ✨ 2. File Upload
# st.sidebar.markdown("📎 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# # Extract text from uploaded file
# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # ✨ 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# # ✨ 4. Tone and Style (shown only for tools that support it)
# show_tone_style = tool_option in ["Paraphraser", "Grammar Checker", "Summarizer"]
# if show_tone_style:
#     tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"])
#     style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"])
# else:
#     tone = style = None

# # ✨ 5. Run Tool
# if st.sidebar.button("Improve Text"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     # if proof_input.strip():
#     #     # System prompts by tool
#     #     tool_prompt_map = {
#     #         "Paraphraser": f"Paraphrase the following text using a {tone.lower()} tone and {style.lower()} style.",
#     #         "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {tone.lower()} and style {style.lower()}.",
#     #         "AI Detector": "Analyze whether this text appears to be written by an AI or a human. Explain your reasoning.",
#     #         "Plagiarism Checker": "Check this text for potential plagiarism. Highlight any copied or suspicious parts.",
#     #         "Summarizer": f"Summarize the following text using a {tone.lower()} tone and {style.lower()} style."
#     #     }

#     #     prompt = tool_prompt_map.get(tool_option, "Improve the following text:")

#     safe_tone = tone.lower() if tone else "neutral"
#     safe_style = style.lower() if style else "default"

#     tool_prompt_map = {
#         "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#         "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {safe_tone} and style {safe_style}.",
#         "AI Detector": "Analyze whether this text appears to be written by an AI or a human. Explain your reasoning.",
#         "Plagiarism Checker": "Check this text for potential plagiarism. Highlight any copied or suspicious parts.",
#         "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#         }
#     prompt = tool_prompt_map.get(tool_option, "Improve the following text:")


#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",
#                     messages=[
#                         {"role": "system", "content": prompt},
#                         {"role": "user", "content": proof_input}
#                     ]
#                 )
#                 tool_output = response['message']['content']
#                 st.session_state["proof_result"] = tool_output

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text for processing.")

# # ✨ 6. Show Output (and optional download)
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;">📝 Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

#     if st.button("Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button("📥 Download", result_bytes, file_name="Improved_text.txt", mime="text/plain")

# st.sidebar.markdown("---")
# st.sidebar.subheader(" AI Proofreader")

# # ✨ 1. Choose Tool First
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # ✨ 2. File Upload
# st.sidebar.markdown("📎 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# # Extract text from uploaded file
# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # ✨ 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# # ✨ 4. Tone and Style (shown only for tools that support it)
# show_tone_style = tool_option in ["Paraphraser", "Grammar Checker", "Summarizer"]
# if show_tone_style:
#     tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"])
#     style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"])
# else:
#     tone = style = None

# # ✨ 5. Run Tool
# if st.sidebar.button("Improve Text"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         # tool_prompt_map = {
#         #     "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#         #     "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {safe_tone} and style {safe_style}.",
#         #     "AI Detector": "Analyze whether this text appears to be written by an AI or a human. Explain your reasoning.",
#         #     "Plagiarism Checker": "Check this text for potential plagiarism. Highlight any copied or suspicious parts.",
#         #     "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#         # }

#         tool_prompt_map = {
#             "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#             "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {safe_tone} and style {safe_style}.",
#             "AI Detector": """Analyze the text and return a JSON response with:
#         {
#             "ai_score": percentage likelihood it's AI-written (0-100),
#             "verdict": short sentence like "Likely AI-generated" or "Likely human-written",
#             "highlight": [list of exact lines or sentences that seem AI-generated]
#         }""",
#             "Plagiarism Checker": """Check this text for plagiarism and return a JSON response with:
#         {
#             "plagiarism_score": percentage of suspected plagiarism (0-100),
#             "plagiarized_lines": [list of exact lines that are likely plagiarized]
#         }""",
#             "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#         }


#         prompt = tool_prompt_map.get(tool_option, "Improve the following text:")

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",
#                     messages=[
#                         {"role": "system", "content": prompt},
#                         {"role": "user", "content": proof_input}
#                     ]
#                 )
#                 # tool_output = response['message']['content']
#                 # st.session_state["proof_result"] = tool_output

#                 import json

#                 raw_output = response['message']['content']

# # Try parsing JSON if AI Detector or Plagiarism Checker
#                 if tool_option in ["AI Detector", "Plagiarism Checker"]:
#                     try:
#                         result_json = json.loads(raw_output)

#                         if tool_option == "AI Detector":
#                             score = result_json.get("ai_score", 0)
#                             verdict = result_json.get("verdict", "No verdict")
#                             highlights = result_json.get("highlight", [])

#                             html = f"""
#                             <h4> AI Detection Score: {score}%</h4>
#                             <p><b>Verdict:</b> {verdict}</p>
#                             <hr>
#                             <h5>🔍 Highlighted AI-like lines:</h5>
#                             """
#                             for line in highlights:
#                                 html += f"<div style='background-color:#ffe6e6; padding:8px; border-radius:5px; margin-bottom:5px;'>{line}</div>"

#                             st.session_state["proof_result"] = html

#                         elif tool_option == "Plagiarism Checker":
#                             score = result_json.get("plagiarism_score", 0)
#                             lines = result_json.get("plagiarized_lines", [])

#                             html = f"""
#                             <h4>📄 Plagiarism Score: {score}%</h4>
#                             <hr>
#                             <h5>🔍 Plagiarized Lines:</h5>
#                             """
#                             for line in lines:
#                                 html += f"<div style='background-color:#ffcccc; padding:8px; border-radius:5px; margin-bottom:5px;'>{line}</div>"

#                             st.session_state["proof_result"] = html

#                     except json.JSONDecodeError:
#                         st.session_state["proof_result"] = f"<p style='color:red;'>❌ Could not parse AI response as JSON. Raw output:</p><pre>{raw_output}</pre>"

#                 else:
#     # For all other tools, just show the raw output
#                     st.session_state["proof_result"] = f"<div>{raw_output}</div>"


#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text for processing.")

# # ✨ 6. Show Output (and optional download)
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;">📝 Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

#     if st.button("Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button(" Download", result_bytes, file_name="Improved_text.txt", mime="text/plain")


#--------------------
# st.sidebar.markdown("---")
# st.sidebar.subheader("📝 AI Proofreader")

# # 1. Choose Tool
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # 2. File Upload
# st.sidebar.markdown("📎 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# # 4. Tone and Style
# show_tone_style = tool_option in ["Paraphraser", "Grammar Checker", "Summarizer"]
# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
# style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# # 5. Improve Button
# if st.sidebar.button("Improve Text"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         tool_prompt_map = {
#             "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#             "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {safe_tone} and style {safe_style}.",
#             "AI Detector": """Analyze the text and return a JSON response with:
# {
#   "ai_score": percentage likelihood it's AI-written (0-100),
#   "verdict": short sentence like "Likely AI-generated" or "Likely human-written",
#   "highlight": [list of exact lines or sentences that seem AI-generated]
# }""",
#             "Plagiarism Checker": """Check this text for plagiarism and return a JSON response with:
# {
#   "plagiarism_score": percentage of suspected plagiarism (0-100),
#   "plagiarized_lines": [list of exact lines that are likely plagiarized]
# }""",
#             "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#         }

#         prompt = tool_prompt_map.get(tool_option, "Improve the following text:")

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",
#                     messages=[
#                         {"role": "system", "content": prompt},
#                         {"role": "user", "content": proof_input}
#                     ]
#                 )

#                 import json
#                 raw_output = response['message']['content']

#                 if tool_option in ["AI Detector", "Plagiarism Checker"]:
#                     try:
#                         result_json = json.loads(raw_output)

#                         if tool_option == "AI Detector":
#                             score = result_json.get("ai_score", 0)
#                             verdict = result_json.get("verdict", "No verdict")
#                             highlights = result_json.get("highlight", [])

#                             html = f"""
#                             <h4>🧠 AI Detection Score: {score}%</h4>
#                             <p><b>Verdict:</b> {verdict}</p>
#                             <hr>
#                             <h5>🔍 Highlighted AI-like lines:</h5>
#                             """
#                             for line in highlights:
#                                 html += f"<div style='background-color:#ffe6e6; padding:8px; border-radius:5px; margin-bottom:5px;'>{line}</div>"

#                             st.session_state["proof_result"] = html

#                         elif tool_option == "Plagiarism Checker":
#                             score = result_json.get("plagiarism_score", 0)
#                             lines = result_json.get("plagiarized_lines", [])

#                             html = f"""
#                             <h4>📄 Plagiarism Score: {score}%</h4>
#                             <hr>
#                             <h5>🔍 Plagiarized Lines:</h5>
#                             """
#                             for line in lines:
#                                 html += f"<div style='background-color:#ffcccc; padding:8px; border-radius:5px; margin-bottom:5px;'>{line}</div>"

#                             st.session_state["proof_result"] = html

#                     except json.JSONDecodeError:
#                         st.session_state["proof_result"] = f"<p style='color:red;'>❌ Could not parse AI response as JSON. Raw output:</p><pre>{raw_output}</pre>"
#                 else:
#                     st.session_state["proof_result"] = f"<div>{raw_output}</div>"

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text for processing.")

# # 6. Output & Download
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;">📝 Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

#     if st.button("Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button("📥 Download", result_bytes, file_name="Improved_text.txt", mime="text/plain")


# from PyPDF2 import PdfReader
# import docx2txt
# import ollama
# import json
# import re
# import streamlit as st

# st.sidebar.markdown("---")
# st.sidebar.subheader("📝 AI Proofreader")

# # 1. Choose Tool
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # 2. File Upload
# st.sidebar.markdown("📌 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")
# # Flatten hard line breaks (like from PDFs or copy-paste)
# if proof_input:
#     proof_input = re.sub(r'\n', ' ', proof_input)  # turn newlines into spaces
#     proof_input = re.sub(r'\s+', ' ', proof_input)  # normalize whitespace

# # 4. Tone and Style
# show_tone_style = tool_option in ["Paraphraser", "Grammar Checker", "Summarizer"]
# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
# style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# # 5. Improve Button
# if st.sidebar.button("Generate Response"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         tool_prompt_map = {
#             "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#             "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {safe_tone} and style {safe_style}.",
#             "AI Detector": """Analyze the text and return a JSON response with:
# {
#   \"ai_score\": percentage likelihood it's AI-written (0-100),
#   \"verdict\": short sentence like \"Likely AI-generated\" or \"Likely human-written\",
#   \"highlight\": [list of exact lines or sentences that seem AI-generated]
# }""",
#             "Plagiarism Checker": """Check this text for plagiarism and return a JSON response with:
# {
#   \"plagiarism_score\": percentage of suspected plagiarism (0-100),
#   \"plagiarized_lines\": [list of exact lines that are likely plagiarized]
# }""",
#             "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#         }

#         prompt = tool_prompt_map.get(tool_option, "Improve the following text:")

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",
#                     messages=[
#                         {"role": "system", "content": prompt},
#                         {"role": "user", "content": proof_input}
#                     ]
#                 )

#                 raw_output = response['message']['content']

#                 if tool_option in ["AI Detector", "Plagiarism Checker"]:
#                     try:
#                         match = re.search(r'\{[\s\S]*?\}', raw_output)
#                         if match:
#                             json_str = match.group(0)
#                             result_json = json.loads(json_str)

#                             #original_lines = proof_input.strip().split("\n")
#                             # original_lines = re.split(r'(?<=[.!?]) +', proof_input.strip())
#                             original_lines = re.split(r'(?<=[.!?]) +', proof_input.strip())

#                             if tool_option == "AI Detector":
#                                 score = result_json.get("ai_score", 0)
#                                 verdict = result_json.get("verdict", "No verdict")
#                                 highlights = set(result_json.get("highlight", []))

#                                 html = f"<h4> AI Detection Score: {score}%</h4><p><b>Verdict:</b> {verdict}</p><hr>"
#                                 html += "<div style='background-color:#ffffff; padding:15px; border-radius:10px;'>"

#                                 for line in original_lines:
#                                     if not line.strip():
#                                         html += "<br>"
#                                     elif line.strip() in highlights:
#                                         html += f"<div style='background-color:#ffe6e6; padding:6px; border-radius:6px; margin-bottom:5px;'>{line}</div>"
#                                     else:
#                                         html += f"<div>{line}</div>"

#                                 html += "</div>"
#                                 st.session_state["proof_result"] = html

#                             elif tool_option == "Plagiarism Checker":
#                                 score = result_json.get("plagiarism_score", 0)
#                                 highlights = set(result_json.get("plagiarized_lines", []))

#                                 html = f"<h4> Plagiarism Score: {score}%</h4><hr>"
#                                 html += "<div style='background-color:#ffffff; padding:15px; border-radius:10px;'>"

#                                 for line in original_lines:
#                                     if not line.strip():
#                                         html += "<br>"
#                                     elif any(highlight.strip() in line for highlight in highlights):
#                                         html += f"<div style='background-color:#ffcccc; padding:6px; border-radius:6px; margin-bottom:5px;'>{line}</div>"
#                                     else:
#                                         html += f"<div>{line}</div>"

#                                 html += "</div>"
#                                 st.session_state["proof_result"] = html
#                         else:
#                             st.session_state["proof_result"] = f"<p style='color:red;'>\u274c No JSON block found in model response:</p><pre>{raw_output}</pre>"

#                     except json.JSONDecodeError:
#                         st.session_state["proof_result"] = f"<p style='color:red;'>\u274c Failed to parse JSON content. Raw output:</p><pre>{raw_output}</pre>"
#                 else:
#                     st.session_state["proof_result"] = f"<div>{raw_output}</div>"

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text for processing.")

# # 6. Show Output & Download Option
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;"> Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

#     if st.button("Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button("\ud83d\udcc5 Download", result_bytes, file_name="Improved_text.txt", mime="text/plain")


# from PyPDF2 import PdfReader
# import docx2txt
# import ollama
# import json
# import re
# import streamlit as st

# st.sidebar.markdown("---")
# st.sidebar.subheader("📝 AI Proofreader")

# # 1. Choose Tool
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # 2. File Upload
# st.sidebar.markdown("📌 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = " ".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# # Pre-clean formatting issues from pasted text
# if proof_input:
#     proof_input = re.sub(r'\n', ' ', proof_input)
#     proof_input = re.sub(r'\s+', ' ', proof_input).strip()

# # 4. Tone and Style
# show_tone_style = tool_option in ["Paraphraser", "Grammar Checker", "Summarizer"]
# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
# style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# # 5. Improve Button
# if st.sidebar.button("Generate Response"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         tool_prompt_map = {
#             "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#             "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {safe_tone} and style {safe_style}.",
#             "AI Detector": """Analyze the text and return a JSON response with:
# {
#   \"ai_score\": percentage likelihood it's AI-written (0-100),
#   \"verdict\": short sentence like \"Likely AI-generated\" or \"Likely human-written\",
#   \"highlight\": [list of exact lines or sentences that seem AI-generated]
# }""",
#             "Plagiarism Checker": """Check this text for plagiarism and return a JSON response with:
# {
#   \"plagiarism_score\": percentage of suspected plagiarism (0-100),
#   \"plagiarized_lines\": [list of exact lines that are likely plagiarized]
# }""",
#             "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#         }

#         prompt = tool_prompt_map.get(tool_option, "Improve the following text:")

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",
#                     messages=[
#                         {"role": "system", "content": prompt},
#                         {"role": "user", "content": proof_input}
#                     ]
#                 )

#                 raw_output = response['message']['content']

#                 if tool_option in ["AI Detector", "Plagiarism Checker"]:
#                     try:
#                         match = re.search(r'\{[\s\S]*?\}', raw_output)
#                         if match:
#                             json_str = match.group(0)
#                             result_json = json.loads(json_str)

#                             # Split by sentence for more accurate highlighting
#                             original_lines = re.split(r'(?<=[.!?]) +', proof_input.strip())

#                             if tool_option == "AI Detector":
#                                 score = result_json.get("ai_score", 0)
#                                 verdict = result_json.get("verdict", "No verdict")
#                                 highlights = set(result_json.get("highlight", []))

#                                 html = f"<h4> AI Detection Score: {score}%</h4><p><b>Verdict:</b> {verdict}</p><hr>"
#                                 html += "<div style='background-color:#ffffff; padding:15px; border-radius:10px;'>"

#                                 for line in original_lines:
#                                     if not line.strip():
#                                         html += "<br>"
#                                     elif any(h.strip() in line for h in highlights):
#                                         html += f"<div style='background-color:#ffe6e6; padding:6px; border-radius:6px; margin-bottom:5px;'>{line}</div>"
#                                     else:
#                                         html += f"<div>{line}</div>"

#                                 html += "</div>"
#                                 st.session_state["proof_result"] = html

#                             elif tool_option == "Plagiarism Checker":
#                                 score = result_json.get("plagiarism_score", 0)
#                                 highlights = set(result_json.get("plagiarized_lines", []))

#                                 html = f"<h4> Plagiarism Score: {score}%</h4><hr>"
#                                 html += "<div style='background-color:#ffffff; padding:15px; border-radius:10px;'>"

#                                 for line in original_lines:
#                                     if not line.strip():
#                                         html += "<br>"
#                                     elif any(h.strip() in line for h in highlights):
#                                         html += f"<div style='background-color:#ffcccc; padding:6px; border-radius:6px; margin-bottom:5px;'>{line}</div>"
#                                     else:
#                                         html += f"<div>{line}</div>"

#                                 html += "</div>"
#                                 st.session_state["proof_result"] = html
#                         else:
#                             st.session_state["proof_result"] = f"<p style='color:red;'>❌ No JSON block found in model response:</p><pre>{raw_output}</pre>"

#                     except json.JSONDecodeError:
#                         st.session_state["proof_result"] = f"<p style='color:red;'>❌ Failed to parse JSON content. Raw output:</p><pre>{raw_output}</pre>"
#                 else:
#                     st.session_state["proof_result"] = f"<div>{raw_output}</div>"

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text for processing.")

# # 6. Show Output & Download Option
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;"> Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

#     if st.button("Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button("📥 Download", result_bytes, file_name="Improved_text.txt", mime="text/plain")

#------------------------working------------------------------------------- Condition---------------------------

# from PyPDF2 import PdfReader
# import docx2txt
# import ollama
# import json
# import re
# import streamlit as st

# st.sidebar.markdown("---")
# st.sidebar.subheader(" AI Proofreader")

# # 1. Choose Tool
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # 2. File Upload
# st.sidebar.markdown(" Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = " ".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# # Pre-clean formatting issues from pasted text
# if proof_input:
#     proof_input = re.sub(r'\n', ' ', proof_input)
#     proof_input = re.sub(r'\s+', ' ', proof_input).strip()

# # 4. Tone and Style
# show_tone_style = tool_option in ["Paraphraser", "Grammar Checker", "Summarizer"]
# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
# style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# # 5. Improve Button
# if st.sidebar.button("Generate Response"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         tool_prompt_map = {
#             "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#             "Grammar Checker": f"Fix grammar and clarity in the following text. Keep tone {safe_tone} and style {safe_style}.",
#             "AI Detector": """Analyze the text and return a JSON response with:
# {
#   \"ai_score\": percentage likelihood it's AI-written (0-100),
#   \"verdict\": short sentence like \"Likely AI-generated\" or \"Likely human-written\",
#   \"highlight\": [list of exact lines or sentences that seem AI-generated]
# }""",
#             "Plagiarism Checker": """Check this text for plagiarism and return a JSON response with:
# {
#   \"plagiarism_score\": percentage of suspected plagiarism (0-100),
#   \"plagiarized_lines\": [list of exact lines that are likely plagiarized]
# }""",
#             "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#         }

#         prompt = tool_prompt_map.get(tool_option, "Improve the following text:")

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",
#                     messages=[
#                         {"role": "system", "content": prompt},
#                         {"role": "user", "content": proof_input}
#                     ]
#                 )

#                 raw_output = response['message']['content']

#                 if tool_option in ["AI Detector", "Plagiarism Checker"]:
#                     try:
#                         match = re.search(r'\{[\s\S]*?\}', raw_output)
#                         if match:
#                             json_str = match.group(0)
#                             result_json = json.loads(json_str)

#                             original_lines = re.split(r'(?<=[.!?]) +', proof_input.strip())

#                             if tool_option == "AI Detector":
#                                 score = result_json.get("ai_score", 0)
#                                 verdict = result_json.get("verdict", "No verdict")
#                                 highlights = result_json.get("highlight", [])

#                                 html = f"<h4> AI Detection Score: {score}%</h4><p><b>Verdict:</b> {verdict}</p><hr>"
#                                 html += "<div style='background-color:#ffffff; padding:15px; border-radius:10px;'>"

#                                 for line in original_lines:
#                                     if not line.strip():
#                                         html += "<br>"
#                                     elif any(h.strip().lower() in line.lower() for h in highlights):
#                                         html += f"<div style='background-color:#ffe6e6; padding:6px; border-radius:6px; margin-bottom:5px;'>{line}</div>"
#                                     else:
#                                         html += f"<div>{line}</div>"

#                                 html += "</div>"
#                                 st.session_state["proof_result"] = html

#                             elif tool_option == "Plagiarism Checker":
#                                 score = result_json.get("plagiarism_score", 0)
#                                 highlights = result_json.get("plagiarized_lines", [])

#                                 html = f"<h4> Plagiarism Score: {score}%</h4><hr>"
#                                 html += "<div style='background-color:#ffffff; padding:15px; border-radius:10px;'>"

#                                 for line in original_lines:
#                                     if not line.strip():
#                                         html += "<br>"
#                                     elif any(h.strip().lower() in line.lower() for h in highlights):
#                                         html += f"<div style='background-color:#ffcccc; padding:6px; border-radius:6px; margin-bottom:5px;'>{line}</div>"
#                                     else:
#                                         html += f"<div>{line}</div>"

#                                 html += "</div>"
#                                 st.session_state["proof_result"] = html
#                         else:
#                             st.session_state["proof_result"] = f"<p style='color:red;'>❌ No JSON block found in model response:</p><pre>{raw_output}</pre>"

#                     except json.JSONDecodeError:
#                         st.session_state["proof_result"] = f"<p style='color:red;'>❌ Failed to parse JSON content. Raw output:</p><pre>{raw_output}</pre>"
#                 else:
#                     st.session_state["proof_result"] = f"<div>{raw_output}</div>"

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text for processing.")

# # 6. Show Output & Download Option
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;"> Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:15px; border-radius:10px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

#     if st.button("Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button("📥 Download", result_bytes, file_name="Improved_text.txt", mime="text/plain")

#--------------------Till Here---------------------------------
# from PyPDF2 import PdfReader
# import ollama
# import json
# import re
# import streamlit as st

# # st.set_page_config(page_title="AI Proofreader", layout="wide")
# st.sidebar.markdown("---")
# st.sidebar.subheader("📚 AI Proofreader")

# # 1. Choose Tool
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # 2. File Upload
# st.sidebar.markdown("📎 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = " ".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             import docx2txt
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# if proof_input:
#     proof_input = re.sub(r'\n+', ' ', proof_input.strip())

# # 4. Tone and Style
# show_tone_style = tool_option in ["Paraphraser", "Grammar Checker", "Summarizer"]
# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
# style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# # 5. Run Tool
# if st.sidebar.button("Generate Response"):
#     st.session_state.pop("ocr_result", None)
#     st.session_state.pop("doc_qa", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         if tool_option == "AI Detector":
#             system_msg = """You are a JSON-only AI content classifier. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
# {
#   \"ai_score\": integer between 0-100,
#   \"verdict\": string,
#   \"highlight\": [sentences likely AI-generated]
# }"""
#             user_msg = proof_input

#         elif tool_option == "Plagiarism Checker":
#             system_msg = """You are a JSON-only plagiarism checker. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
# {
#   \"plagiarism_score\": integer between 0-100,
#   \"plagiarized_lines\": [sentences likely plagiarized]
# }"""
#             user_msg = proof_input

#         else:
#             prompt = {
#                 "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#                 "Grammar Checker": f"Fix grammar and clarity in the following text for US Curriculumn following NGSS Guidelines. Keep tone {safe_tone} and style {safe_style}.",
#                 "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#             }.get(tool_option)
#             system_msg = None
#             user_msg = prompt

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 messages = []
#                 if system_msg:
#                     messages.append({"role": "system", "content": system_msg})
#                 messages.append({"role": "user", "content": user_msg})

#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",
#                     messages=messages
#                 )
#                 raw_output = response['message']['content']

#                 # Updated regex (no ?R, works with Python)
#                 match = re.search(r'\{(?:[^{}"\\]|\\.|"(?:\\.|[^"\\])*")*\}', raw_output, re.DOTALL)

#                 if tool_option in ["AI Detector", "Plagiarism Checker"] and match:
#                     try:
#                         result_json = json.loads(match.group(0))
#                         lines = re.split(r'(?<=[.!?])\s+', proof_input.strip())
#                         highlights = result_json.get("highlight" if tool_option == "AI Detector" else "plagiarized_lines", [])
#                         score = result_json.get("ai_score" if tool_option == "AI Detector" else "plagiarism_score", 0)
#                         verdict = result_json.get("verdict", "") if tool_option == "AI Detector" else ""
#                         highlight_color = "#ffe6e6" if tool_option == "AI Detector" else "#ffcccc"

#                         html = f"""
#                         <div style='padding: 20px; background-color: #fff; border-radius: 12px;'>
#                             <h3 style='margin-top:0;'>{'🧠 AI Detection Score' if tool_option == 'AI Detector' else '📄 Plagiarism Score'}: {score}%</h3>
#                             {f"<p><b>Verdict:</b> {verdict}</p>" if verdict else ''}
#                             <hr>
#                         """
#                         for line in lines:
#                             line_clean = line.strip()
#                             if any(h.strip().lower() in line_clean.lower() for h in highlights):
#                                 html += f"<div style='background-color:{highlight_color}; padding:10px; margin-bottom:8px; border-radius:8px;'>{line}</div>"
#                             else:
#                                 html += f"<p style='margin: 5px 0;'>{line}</p>"
#                         html += "</div>"
#                         st.session_state["proof_result"] = html

#                     except Exception as e:
#                         st.session_state["proof_result"] = f"<p style='color:red;'>❌ Error parsing JSON:</p><pre>{match.group(0)}</pre><p>{e}</p>"
#                 else:
#                     st.warning("⚠️ No JSON found in model response. Showing fallback explanation.")
#                     st.session_state["proof_result"] = f"<pre style='background:#f8f8f8;padding:10px;border-radius:6px;'>{raw_output}</pre>"

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text.")

# # 6. Show Output & Download Option
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;">📋 Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:20px; border-radius:12px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )

#     if st.button("📥 Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button("📥 Save", result_bytes, file_name="Improved_text.txt", mime="text/plain")





#--------------------17th April---------------------------------------------------------


# from PyPDF2 import PdfReader
# import ollama
# import json
# import re
# import streamlit as st
# import difflib
# from docx import Document
# from docx.enum.text import WD_COLOR_INDEX

# # Streamlit setup
# st.sidebar.markdown("---")
# st.sidebar.subheader("📚 AI Proofreader")

# # 1. Choose Tool
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # 2. File Upload
# st.sidebar.markdown("📎 Upload a file (PDF, Word, or TXT):")
# uploaded_proof_file = st.sidebar.file_uploader("", type=["pdf", "docx", "txt"])

# if uploaded_proof_file:
#     try:
#         extracted_text = ""
#         if uploaded_proof_file.type == "application/pdf":
#             pdf = PdfReader(uploaded_proof_file)
#             extracted_text = " ".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         elif uploaded_proof_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             import docx2txt
#             extracted_text = docx2txt.process(uploaded_proof_file)
#         elif uploaded_proof_file.type == "text/plain":
#             extracted_text = uploaded_proof_file.read().decode("utf-8")

#         if extracted_text.strip():
#             if st.sidebar.button("Use Extracted Text"):
#                 st.session_state["proof_input"] = extracted_text
#                 st.sidebar.success("Text loaded into editor below!")
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # 3. Enter/Edit Text
# proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# if proof_input:
#     proof_input = re.sub(r'\n+', ' ', proof_input.strip())

# # 4. Tone and Style (skip for Grammar Checker)
# show_tone_style = tool_option in ["Paraphraser", "Summarizer"]
# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
# style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# # Helper: Highlight grammar changes in Word
# def add_highlighted_text(paragraph, original, corrected):
#     matcher = difflib.SequenceMatcher(None, original.split(), corrected.split())
#     for tag, i1, i2, j1, j2 in matcher.get_opcodes():
#         if tag == "equal":
#             paragraph.add_run(" " + " ".join(original.split()[i1:i2]))
#         elif tag in ("replace", "insert"):
#             run = paragraph.add_run(" " + " ".join(corrected.split()[j1:j2]))
#             run.font.highlight_color = WD_COLOR_INDEX.YELLOW
#         elif tag == "delete":
#             continue

# # 5. Run Tool
# if st.sidebar.button("Generate Response"):
#     st.session_state.pop("proof_result", None)
#     st.session_state.pop("proof_doc_path", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         if tool_option == "AI Detector":
#             system_msg = """You are a JSON-only AI content classifier. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
# {
#   "ai_score": integer between 0-100,
#   "verdict": string,
#   "highlight": [sentences likely AI-generated]
# }"""
#             user_msg = proof_input

#         elif tool_option == "Plagiarism Checker":
#             system_msg = """You are a JSON-only plagiarism checker. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
# {
#   "plagiarism_score": integer between 0-100,
#   "plagiarized_lines": [sentences likely plagiarized]
# }"""
#             user_msg = proof_input

#         elif tool_option == "Grammar Checker":
#             system_msg = None
#             user_msg = f"Fix grammar and clarity in the following text for US Curriculum following NGSS Guidelines. Do not change tone or meaning. Do not paraphrase.\n\n{proof_input}"

#         else:
#             prompt = {
#                 "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#                 "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#             }.get(tool_option)
#             system_msg = None
#             user_msg = prompt + "\n\n" + proof_input if prompt else proof_input

#         with st.spinner(f"Running {tool_option}..."):
#             try:
#                 messages = []
#                 if system_msg:
#                     messages.append({"role": "system", "content": system_msg})
#                 messages.append({"role": "user", "content": user_msg})

#                 response = ollama.chat(
#                     model="llama3.2-vision:latest",  # Swap to phi3 or mistral for speed if needed
#                     messages=messages
#                 )
#                 raw_output = response['message']['content'].strip()

#                 if tool_option == "Grammar Checker":
#                     # ✅ Save Word doc locally (not in /mnt/data)
#                     doc = Document()
#                     doc.add_heading("Grammar-Checked Document", level=1)
#                     para = doc.add_paragraph()
#                     add_highlighted_text(para, proof_input, raw_output)
#                     word_path = "grammar_checked.docx"
#                     doc.save(word_path)
#                     st.session_state["proof_doc_path"] = word_path
#                     st.success("✅ Grammar check complete! Download your Word doc below.")

#                 elif tool_option in ["AI Detector", "Plagiarism Checker"]:
#                     match = re.search(r'\{(?:[^{}"\\]|\\.|"(?:\\.|[^"\\])*")*\}', raw_output, re.DOTALL)
#                     if match:
#                         try:
#                             result_json = json.loads(match.group(0))
#                             lines = re.split(r'(?<=[.!?])\s+', proof_input.strip())
#                             highlights = result_json.get("highlight" if tool_option == "AI Detector" else "plagiarized_lines", [])
#                             score = result_json.get("ai_score" if tool_option == "AI Detector" else "plagiarism_score", 0)
#                             verdict = result_json.get("verdict", "") if tool_option == "AI Detector" else ""
#                             highlight_color = "#ffe6e6" if tool_option == "AI Detector" else "#ffcccc"

#                             html = f"""
#                             <div style='padding: 20px; background-color: #fff; border-radius: 12px;'>
#                                 <h3 style='margin-top:0;'>{'🧠 AI Detection Score' if tool_option == 'AI Detector' else '📄 Plagiarism Score'}: {score}%</h3>
#                                 {f"<p><b>Verdict:</b> {verdict}</p>" if verdict else ''}
#                                 <hr>
#                             """
#                             for line in lines:
#                                 line_clean = line.strip()
#                                 if any(h.strip().lower() in line_clean.lower() for h in highlights):
#                                     html += f"<div style='background-color:{highlight_color}; padding:10px; margin-bottom:8px; border-radius:8px;'>{line}</div>"
#                                 else:
#                                     html += f"<p style='margin: 5px 0;'>{line}</p>"
#                             html += "</div>"
#                             st.session_state["proof_result"] = html
#                         except Exception as e:
#                             st.session_state["proof_result"] = f"<p style='color:red;'>❌ Error parsing JSON:</p><pre>{match.group(0)}</pre><p>{e}</p>"
#                     else:
#                         st.warning("⚠️ No JSON found in model response. Showing fallback explanation.")
#                         st.session_state["proof_result"] = f"<pre style='background:#f8f8f8;padding:10px;border-radius:6px;'>{raw_output}</pre>"
#                 else:
#                     st.session_state["proof_result"] = f"<pre style='background:#f8f8f8;padding:10px;border-radius:6px;'>{raw_output}</pre>"

#             except Exception as e:
#                 st.sidebar.error(f"Error: {str(e)}")
#     else:
#         st.sidebar.warning("Please enter some text.")

# # 6. Show Output or Word Download
# if "proof_result" in st.session_state:
#     st.markdown('<h3 style="color: white;">📋 Output:</h3>', unsafe_allow_html=True)
#     st.markdown(
#         f"<div style='background-color:#ffffff; padding:20px; border-radius:12px; color: black;'>{st.session_state['proof_result']}</div>",
#         unsafe_allow_html=True
#     )
#     if st.button("📥 Download Result as .txt"):
#         result_text = st.session_state['proof_result']
#         result_bytes = result_text.encode('utf-8')
#         st.download_button("📥 Save", result_bytes, file_name="Improved_text.txt", mime="text/plain")

# if "proof_doc_path" in st.session_state:
#     with open(st.session_state["proof_doc_path"], "rb") as file:
#         st.download_button(
#             label="📥 Download Grammar-Checked Word Document",
#             data=file,
#             file_name="Grammar_Checked_Document.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

#----------------------17th April-------------------------------

# import streamlit as st
# import pdfplumber
# import ollama
# import json
# import re
# import difflib
# from docx import Document
# from docx.enum.text import WD_COLOR_INDEX

# # st.set_page_config(page_title="AI Proofreader", layout="wide")
# st.sidebar.subheader("📚 AI Proofreader")

# # 1. Choose Tool
# tool_option = st.sidebar.selectbox(
#     "Choose a tool:",
#     ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
# )

# # 2. File Upload
# uploaded_file = st.sidebar.file_uploader("📎 Upload a file (PDF, Word, or TXT):", type=["pdf", "docx", "txt"])

# # Store extracted content
# extracted_text = ""
# layout_paragraphs = []

# if uploaded_file:
#     try:
#         if uploaded_file.type == "application/pdf":
#             with pdfplumber.open(uploaded_file) as pdf:
#                 for page in pdf.pages:
#                     text = page.extract_text()
#                     if text:
#                         lines = text.split('\n')
#                         layout_paragraphs.extend([line.strip() for line in lines if line.strip()])
#             extracted_text = "\n".join(layout_paragraphs)
#         elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             import docx2txt
#             extracted_text = docx2txt.process(uploaded_file)
#         elif uploaded_file.type == "text/plain":
#             extracted_text = uploaded_file.read().decode("utf-8")

#         if extracted_text.strip():
#             st.session_state["proof_input"] = extracted_text.strip()
#         else:
#             st.sidebar.warning("No text found in the uploaded file.")
#     except Exception as e:
#         st.sidebar.error(f"Error reading file: {str(e)}")

# # 3. Editable input for non-PDF tools
# proof_input = st.session_state.get("proof_input", "")
# if tool_option != "Grammar Checker":
#     proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# # 4. Tone and Style
# show_tone_style = tool_option in ["Paraphraser", "Summarizer"]
# tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
# style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# # Highlight differences in Word
# def add_highlighted_diff(paragraph, original, corrected):
#     matcher = difflib.SequenceMatcher(None, original.split(), corrected.split())
#     for tag, i1, i2, j1, j2 in matcher.get_opcodes():
#         if tag == "equal":
#             paragraph.add_run(" " + " ".join(original.split()[i1:i2]))
#         elif tag in ("replace", "insert"):
#             run = paragraph.add_run(" " + " ".join(corrected.split()[j1:j2]))
#             run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# # 5. Generate response
# if st.sidebar.button("Generate Response"):
#     st.session_state.pop("proof_result", None)
#     st.session_state.pop("proof_doc_path", None)

#     if proof_input.strip():
#         safe_tone = tone.lower() if tone else "neutral"
#         safe_style = style.lower() if style else "default"

#         if tool_option == "Grammar Checker":
#             with st.spinner("Running Grammar Checker with layout preserved..."):
#                 doc = Document()
#                 doc.add_heading("Grammar-Checked Document", level=1)
#                 for para_text in layout_paragraphs:
#                     user_msg = f"Fix grammar only in the following text for the US Curriculum, following NGSS Guidelines. Do not paraphrase or alter the wording.\n\n{para_text}"
#                     response = ollama.chat(
#                         model="phi4:latest",  # use a lightweight model
#                         messages=[{"role": "user", "content": user_msg}]
#                     )
#                     fixed = response["message"]["content"].strip()
#                     para = doc.add_paragraph()
#                     add_highlighted_diff(para, para_text, fixed)

#                 output_path = "grammar_checked_layout.docx"
#                 doc.save(output_path)
#                 st.session_state["proof_doc_path"] = output_path
#                 st.success("✅ Grammar check complete! Download below.")

#         else:
#             if tool_option == "AI Detector":
#                 system_msg = """You are a JSON-only AI content classifier. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
# {
#   "ai_score": integer between 0-100,
#   "verdict": string,
#   "highlight": [sentences likely AI-generated]
# }"""
#                 user_msg = proof_input

#             elif tool_option == "Plagiarism Checker":
#                 system_msg = """You are a JSON-only plagiarism checker. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
# {
#   "plagiarism_score": integer between 0-100,
#   "plagiarized_lines": [sentences likely plagiarized]
# }"""
#                 user_msg = proof_input

#             else:
#                 prompt = {
#                     "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
#                     "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
#                 }.get(tool_option)
#                 system_msg = None
#                 user_msg = prompt + "\n\n" + proof_input if prompt else proof_input

#             messages = []
#             if system_msg:
#                 messages.append({"role": "system", "content": system_msg})
#             messages.append({"role": "user", "content": user_msg})

#             with st.spinner(f"Running {tool_option}..."):
#                 try:
#                     response = ollama.chat(
#                         model="llama3.2-vision:latest",
#                         messages=messages
#                     )
#                     raw_output = response['message']['content'].strip()
#                     match = re.search(r'\{(?:[^{}"\\]|\\.|"(?:\\.|[^"\\])*")*\}', raw_output, re.DOTALL)

#                     if tool_option in ["AI Detector", "Plagiarism Checker"] and match:
#                         result_json = json.loads(match.group(0))
#                         lines = re.split(r'(?<=[.!?])\s+', proof_input.strip())
#                         highlights = result_json.get("highlight" if tool_option == "AI Detector" else "plagiarized_lines", [])
#                         score = result_json.get("ai_score" if tool_option == "AI Detector" else "plagiarism_score", 0)
#                         verdict = result_json.get("verdict", "") if tool_option == "AI Detector" else ""
#                         highlight_color = "#ffe6e6" if tool_option == "AI Detector" else "#ffcccc"

#                         html = f"""
#                         <div style='padding: 20px; background-color: #fff; border-radius: 12px;'>
#                             <h3>{'🧠 AI Detection Score' if tool_option == 'AI Detector' else '📄 Plagiarism Score'}: {score}%</h3>
#                             {f"<p><b>Verdict:</b> {verdict}</p>" if verdict else ''}
#                             <hr>
#                         """
#                         for line in lines:
#                             line_clean = line.strip()
#                             if any(h.strip().lower() in line_clean.lower() for h in highlights):
#                                 html += f"<div style='background-color:{highlight_color}; padding:10px; margin-bottom:8px; border-radius:8px;'>{line}</div>"
#                             else:
#                                 html += f"<p style='margin: 5px 0;'>{line}</p>"
#                         html += "</div>"
#                         st.session_state["proof_result"] = html
#                     else:
#                         st.session_state["proof_result"] = f"<pre style='background:#f8f8f8;padding:10px;border-radius:6px;'>{raw_output}</pre>"

#                 except Exception as e:
#                     st.error(f"Error: {str(e)}")

# # 6. Show Output & Download
# if "proof_result" in st.session_state:
#     st.markdown("### 📋 Output")
#     st.markdown(st.session_state["proof_result"], unsafe_allow_html=True)

# if "proof_doc_path" in st.session_state:
#     with open(st.session_state["proof_doc_path"], "rb") as f:
#         st.download_button(
#             label="📥 Download Grammar-Checked Word Doc",
#             data=f,
#             file_name="Grammar_Checked_Document.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

#--------------------------------------------------------------------

import streamlit as st
import pdfplumber
import ollama
import json
import re
import difflib
import uuid
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os
import time

st.sidebar.subheader("📚 AI Proofreader")

# Clean up old Word files (>1 hour)
os.makedirs("output_docs", exist_ok=True)
for file in os.listdir("output_docs"):
    path = os.path.join("output_docs", file)
    if file.endswith(".docx") and time.time() - os.path.getmtime(path) > 3600:
        os.remove(path)

# 1. Choose Tool
tool_option = st.sidebar.selectbox(
    "Choose a tool:",
    ["Paraphraser", "Grammar Checker", "AI Detector", "Plagiarism Checker", "Summarizer"]
)

# 2. File Upload
uploaded_file = st.sidebar.file_uploader("📎 Upload a file (PDF, Word, or TXT):", type=["pdf", "docx", "txt"])

# Extract text and preserve layout
extracted_text = ""
layout_paragraphs = []

if uploaded_file:
    try:
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')
                        layout_paragraphs.extend([line.strip() for line in lines if line.strip()])
            extracted_text = "\n".join(layout_paragraphs)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            import docx2txt
            extracted_text = docx2txt.process(uploaded_file)
        elif uploaded_file.type == "text/plain":
            extracted_text = uploaded_file.read().decode("utf-8")

        if extracted_text.strip():
            st.session_state["proof_input"] = extracted_text.strip()
        else:
            st.sidebar.warning("No text found in the uploaded file.")
    except Exception as e:
        st.sidebar.error(f"Error reading file: {str(e)}")

# 3. Text input area (only for non-PDF tools)
proof_input = st.session_state.get("proof_input", "")
if tool_option != "Grammar Checker":
    proof_input = st.sidebar.text_area("Enter text to improve:", height=150, key="proof_input")

# 4. Tone and Style
show_tone_style = tool_option in ["Paraphraser", "Summarizer"]
tone = st.sidebar.selectbox("Tone", ["Neutral", "Formal", "Informal"]) if show_tone_style else None
style = st.sidebar.selectbox("Style", ["Simplify", "Make Concise", "Expand"]) if show_tone_style else None

# Highlight differences in Word doc
def add_highlighted_diff(paragraph, original, corrected):
    matcher = difflib.SequenceMatcher(None, original.split(), corrected.split())
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            paragraph.add_run(" " + " ".join(original.split()[i1:i2]))
        elif tag in ("replace", "insert"):
            run = paragraph.add_run(" " + " ".join(corrected.split()[j1:j2]))
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# 5. Run Tool
if st.sidebar.button("Generate Response"):
    st.session_state.pop("proof_result", None)
    st.session_state.pop("proof_doc_path", None)

    if proof_input.strip():
        safe_tone = tone.lower() if tone else "neutral"
        safe_style = style.lower() if style else "default"

        if tool_option == "Grammar Checker":
            with st.spinner("Running Grammar Checker with layout preserved..."):
                doc = Document()
                doc.add_heading("Grammar-Checked Document", level=1)

                for para_text in layout_paragraphs:
#                     user_msg = f"""Fix only grammar and punctuation errors in the following text.

# - Do NOT paraphrase, reword, expand, summarize, simplify, or explain anything.
# - Do NOT change sentence structure or tone.
# - Keep capitalization, line breaks, and phrasing exactly the same.
# - Output only the corrected version of the text. Nothing else.

# TEXT:
# \"\"\"
# {para_text}
# \"\"\"
# """
                    user_msg = f"""Fix only grammar and punctuation errors in the following text.

- Do NOT paraphrase, simplify, reword, summarize, or explain anything.
- Do NOT change sentence structure or tone.
- Do NOT add or remove formatting characters like quotes, triple backticks, asterisks, or Markdown symbols.
- Keep capitalization, line breaks, and phrasing exactly the same.
- Output ONLY the corrected version of the text. No extra comments or formatting.

TEXT:
{para_text}
"""


                    response = ollama.chat(
                        model="phi4:latest",
                        messages=[{"role": "user", "content": user_msg}]
                    )
                    fixed = response["message"]["content"].strip()
                    para = doc.add_paragraph()
                    add_highlighted_diff(para, para_text, fixed)

                # Save with unique filename
                output_path = os.path.join("output_docs", f"grammar_checked_{uuid.uuid4().hex}.docx")
                doc.save(output_path)
                st.session_state["proof_doc_path"] = output_path
                st.success("✅ Grammar check complete! Download below.")

        else:
            if tool_option == "AI Detector":
                system_msg = """You are a JSON-only AI content classifier. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
{
  "ai_score": integer between 0-100,
  "verdict": string,
  "highlight": [sentences likely AI-generated]
}"""
                user_msg = proof_input

            elif tool_option == "Plagiarism Checker":
                system_msg = """You are a JSON-only plagiarism checker. You MUST return only valid JSON. Do NOT include explanations or markdown. The format is:
{
  "plagiarism_score": integer between 0-100,
  "plagiarized_lines": [sentences likely plagiarized]
}"""
                user_msg = proof_input

            else:
                prompt = {
                    "Paraphraser": f"Paraphrase the following text using a {safe_tone} tone and {safe_style} style.",
                    "Summarizer": f"Summarize the following text using a {safe_tone} tone and {safe_style} style."
                }.get(tool_option)
                system_msg = None
                user_msg = prompt + "\n\n" + proof_input if prompt else proof_input

            messages = []
            if system_msg:
                messages.append({"role": "system", "content": system_msg})
            messages.append({"role": "user", "content": user_msg})

            with st.spinner(f"Running {tool_option}..."):
                try:
                    response = ollama.chat(
                        model="llama3.2-vision:latest",
                        messages=messages
                    )
                    raw_output = response['message']['content'].strip()
                    match = re.search(r'\{(?:[^{}"\\]|\\.|"(?:\\.|[^"\\])*")*\}', raw_output, re.DOTALL)

                    if tool_option in ["AI Detector", "Plagiarism Checker"] and match:
                        result_json = json.loads(match.group(0))
                        lines = re.split(r'(?<=[.!?])\s+', proof_input.strip())
                        highlights = result_json.get("highlight" if tool_option == "AI Detector" else "plagiarized_lines", [])
                        score = result_json.get("ai_score" if tool_option == "AI Detector" else "plagiarism_score", 0)
                        verdict = result_json.get("verdict", "") if tool_option == "AI Detector" else ""
                        highlight_color = "#ffe6e6" if tool_option == "AI Detector" else "#ffcccc"

                        html = f"""
                        <div style='padding: 20px; background-color: #fff; border-radius: 12px;'>
                            <h3>{'🧠 AI Detection Score' if tool_option == 'AI Detector' else '📄 Plagiarism Score'}: {score}%</h3>
                            {f"<p><b>Verdict:</b> {verdict}</p>" if verdict else ''}
                            <hr>
                        """
                        for line in lines:
                            line_clean = line.strip()
                            if any(h.strip().lower() in line_clean.lower() for h in highlights):
                                html += f"<div style='background-color:{highlight_color}; padding:10px; margin-bottom:8px; border-radius:8px;'>{line}</div>"
                            else:
                                html += f"<p style='margin: 5px 0;'>{line}</p>"
                        html += "</div>"
                        st.session_state["proof_result"] = html
                    else:
                        st.session_state["proof_result"] = f"<pre style='background:#f8f8f8;padding:10px;border-radius:6px;'>{raw_output}</pre>"

                except Exception as e:
                    st.error(f"Error: {str(e)}")

# 6. Output Section
if "proof_result" in st.session_state:
    st.markdown("### 📋 Output")
    st.markdown(st.session_state["proof_result"], unsafe_allow_html=True)

if "proof_doc_path" in st.session_state:
    with open(st.session_state["proof_doc_path"], "rb") as f:
        st.download_button(
            label="📥 Download Grammar-Checked Word Doc",
            data=f,
            file_name="Grammar_Checked_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )













# # Insert this additional feature into the original code
# st.sidebar.header("Advanced Features")
# # File upload for advanced processing
# uploaded_file = st.sidebar.file_uploader("Upload a file (PDF, Word, or Text)", type=["pdf", "docx", "txt"])

# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         extracted_content = extract_text_from_pdf(uploaded_file)
#     elif uploaded_file.type.endswith("wordprocessingml.document"):
#         extracted_content = extract_text_from_word(uploaded_file)
#     else:
#         extracted_content = uploaded_file.read().decode("utf-8")

#     if extracted_content:
#         st.sidebar.markdown("### Extracted Content Preview")
#         st.sidebar.text(extracted_content[:500])  # Preview the first 500 characters
#     else:
#         st.sidebar.warning("No content could be extracted from the uploaded file.")

# # Document analysis
# user_advanced_prompt = st.sidebar.text_area("Enter a task for the document:")
# if st.sidebar.button("Analyze Document"):
#     if uploaded_file and extracted_content:
#         analysis_result = analyze_chunks_together(extracted_content, user_advanced_prompt)
#         st.sidebar.markdown("### Analysis Result")
#         st.sidebar.text(analysis_result)
#     else:
#         st.sidebar.warning("Upload a file and ensure text is extracted before analysis.")

# # Free-form chat
# free_chat_input = st.sidebar.text_area("Enter your query for free-form chat:")
# if st.sidebar.button("Submit Query"):
#     if free_chat_input:
#         free_chat_result = generate_openai_response(free_chat_input, "gpt-4", openai.api_key, temperature, max_tokens)
#         st.sidebar.markdown("### Free-Form Chat Result")
#         st.sidebar.text(free_chat_result)
#     else:
#         st.sidebar.warning("Please enter a query to process.")





# import os
# import base64
# import io
# import streamlit as st
# from dotenv import load_dotenv
# from PIL import Image, UnidentifiedImageError
# from pdf2image import convert_from_bytes
# from textstat import flesch_reading_ease
# from transformers import pipeline, AutoModelForCausalLM, AutoTokenizer
# import torch
# import openai

# # Optional: Auto-install docx2txt if missing
# try:
#     import docx2txt
# except ImportError:
#     import subprocess, sys
#     subprocess.check_call([sys.executable, "-m", "pip", "install", "docx2txt"])
#     import docx2txt

# # ------------------ ENV SETUP ------------------
# dotenv_path = os.path.join(os.path.dirname(__file__), ".env")
# load_dotenv(dotenv_path)
# openai.api_key = os.getenv("OPENAI_API_KEY") or "your-api-key-here"

# # ------------------ STYLING ------------------
# def set_background(img_path):
#     if not os.path.exists(img_path):
#         st.warning("Background image not found.")
#         return
#     with open(img_path, 'rb') as f:
#         base64_img = base64.b64encode(f.read()).decode()
#     page_bg_img = f'''
#     <style>
#         .stApp {{
#             background-image: url("data:image/png;base64,{base64_img}");
#             background-size: cover;
#             background-repeat: no-repeat;
#             background-attachment: fixed;
#         }}
#         label, h1 {{ color: black; }}
#     </style>
#     '''
#     st.markdown(page_bg_img, unsafe_allow_html=True)

# # ------------------ HELPERS ------------------
# def get_llm_response(prompt, model):
#     if model.startswith("gpt"):
#         try:
#             response = openai.ChatCompletion.create(
#                 model=model,
#                 messages=[{"role": "user", "content": prompt}],
#                 temperature=st.session_state.temperature,
#                 max_tokens=st.session_state.max_tokens
#             )
#             return response["choices"][0]["message"]["content"]
#         except Exception as e:
#             return f"OpenAI Error: {str(e)}"
#     elif model == "open-llama":
#         try:
#             model_id = "openlm-research/open_llama_7b"
#             tokenizer = AutoTokenizer.from_pretrained(model_id)
#             model = AutoModelForCausalLM.from_pretrained(model_id)
#             device = 0 if torch.cuda.is_available() else -1
#             generator = pipeline("text-generation", model=model, tokenizer=tokenizer, device=device)
#             response = generator(prompt, max_length=st.session_state.max_tokens, temperature=st.session_state.temperature)
#             return response[0]['generated_text']
#         except Exception as e:
#             return f"LLaMA Error: {str(e)}"
#     return "Unsupported model."

# def simplify_text(text):
#     return text  # Add simplification logic if needed

# def adjust_text_for_grade(text, level):
#     scores = {"K-2": 90, "3-5": 80, "6-8": 70, "9-12": 60}
#     target = scores.get(level, 50)
#     while flesch_reading_ease(text) < target:
#         text = simplify_text(text)
#     return text

# # ------------------ SIDEBAR ------------------
# st.sidebar.title("Settings")
# st.session_state.temperature = st.sidebar.slider("Temperature", 0.0, 1.0, 0.7)
# st.session_state.max_tokens = st.sidebar.slider("Max Tokens", 10, 2000, 500)
# st.session_state.user_role = st.sidebar.selectbox("Are you an:", ["Instructor", "Learner"])
# st.session_state.grade_level = st.sidebar.selectbox("Select Grade Level:", ["K-2", "3-5", "6-8", "9-12"])

# model_mapping = {
#     "GPT-4": "gpt-4",
#     "GPT-4 Turbo": "gpt-4-turbo",
#     "GPT-4o": "gpt-4o",
#     "OpenLLaMA": "open-llama"
# }
# model_choice = st.sidebar.selectbox("Select an AI Model", list(model_mapping.keys()))
# llm_model = model_mapping[model_choice]

# # ------------------ MAIN UI ------------------
# st.title("🦙 QBS OCR + AI Tool")
# set_background("C:\\Users\\prateek.kumar\\Desktop\\Langchain Project\\QBS (4).jpg")

# st.markdown("### Upload a File for OCR/Text Analysis")
# uploaded_file = st.file_uploader("Upload a file", type=['png', 'jpg', 'jpeg', 'pdf', 'docx', 'txt'])

# if uploaded_file:
#     all_text = ""
#     if uploaded_file.type.startswith("image"):
#         image = Image.open(uploaded_file)
#         st.image(image, caption="Uploaded Image")
#         if st.button("Extract Text 🔍"):
#             st.session_state.ocr_result = "[Mocked response from vision model for image]"

#     elif uploaded_file.type == "application/pdf":
#         images = convert_from_bytes(uploaded_file.getvalue())
#         if st.button("Extract Text 🔍"):
#             all_text = ""
#             for i, img in enumerate(images):
#                 buf = io.BytesIO()
#                 img.save(buf, format="PNG")
#                 buf.seek(0)
#                 all_text += f"## Page {i+1}\n[Mocked response from vision model]\n\n---\n"
#             st.session_state.ocr_result = all_text

#     elif uploaded_file.type.endswith("document"):
#         if st.button("Extract Text 🔍"):
#             raw_text = docx2txt.process(uploaded_file)
#             st.session_state.ocr_result = f"[Formatted content]\n\n{raw_text}"

#     elif uploaded_file.type == "text/plain":
#         if st.button("Extract Text 🔍"):
#             raw_text = uploaded_file.read().decode("utf-8")
#             st.session_state.ocr_result = f"[Formatted text]\n\n{raw_text}"

# if 'ocr_result' in st.session_state:
#     st.markdown(st.session_state['ocr_result'])

# # ------------------ PROMPT TASKS ------------------
# st.markdown("---")
# st.markdown(f'<h1 class="welcome-text">Welcome {st.session_state.user_role}!</h1>', unsafe_allow_html=True)

# user_prompt = st.text_area("Enter your input:")
# task_type = st.selectbox("Select Task:", ["Grammar Check", "Solve Math Problem", "Generate Image", "Generate Content"])

# if st.button("Process Task") and user_prompt:
#     prompt_map = {
#         "Grammar Check": f"Fix grammar and improve clarity:\n\n{user_prompt}",
#         "Solve Math Problem": f"Solve this step-by-step:\n\n{user_prompt}",
#         "Generate Image": f"Image prompt: {user_prompt}",
#         "Generate Content": user_prompt
#     }
#     result = get_llm_response(prompt_map[task_type], llm_model)
#     st.markdown("### Result:")
#     st.markdown(f"<div class='response-box'>{result}</div>", unsafe_allow_html=True)
